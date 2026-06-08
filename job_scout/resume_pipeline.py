from __future__ import annotations

import json
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

from job_scout.config import MAIN_RESUME_DOCX, TRANSCRIPT_PROFILE_JSON
from job_scout.resume_source import ResumeSource, load_resume_source


SECTION_HEADINGS = [
    "Summary",
    "Education",
    "Technical Skills",
    "Projects",
    "Experience",
    "Leadership & Activities",
]

APPROVED_GPA_LINE = (
    "GPA: 3.47 Overall | 3.65 CS Coursework | "
    "Dean's List: Fall 2023, Spring 2025, Fall 2025"
)

STATIC_FORBIDDEN_VISIBLE_TEXT = [
    "Tailored for",
    "transcript-backed strengths",
    "Computer Science undergraduate",
    "GPA 3.46",
    "Dept GPA",
    "Major GPA",
    "Ayush Yadav Resume (Dec 2025)",
]


@dataclass(frozen=True)
class ResumeArtifacts:
    docx_path: Path
    pdf_path: Path
    evidence_path: Path
    validation_path: Path


def _slug(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "-", value).strip("-")
    return cleaned or "Resume"


def _artifact_company(value: str) -> str:
    cleaned = re.sub(r"[{}\\/]+", "", str(value)).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned or "Company"


def _paragraph_texts(doc: Document) -> list[str]:
    texts: list[str] = []
    for paragraph in doc.paragraphs:
        for part in paragraph.text.splitlines():
            text = part.strip()
            if text:
                texts.append(text)
    return texts


def _append_unique(texts: list[str], seen: set[str], text: str) -> None:
    cleaned = re.sub(r"\s+", " ", text.strip())
    if cleaned and cleaned not in seen:
        texts.append(cleaned)
        seen.add(cleaned)


def _iter_table_texts(table: Table) -> list[str]:
    texts: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for part in paragraph.text.splitlines():
                    text = part.strip()
                    if text:
                        texts.append(text)
            for nested_table in cell.tables:
                texts.extend(_iter_table_texts(nested_table))
    return texts


def _all_docx_texts(doc: Document) -> list[str]:
    texts: list[str] = []
    seen: set[str] = set()
    for header_attr in ("first_page_header", "header", "even_page_header"):
        for section in doc.sections:
            header = getattr(section, header_attr)
            for paragraph in header.paragraphs:
                for part in paragraph.text.splitlines():
                    _append_unique(texts, seen, part)
            for table in header.tables:
                for text in _iter_table_texts(table):
                    _append_unique(texts, seen, text)
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            for part in Paragraph(child, doc).text.splitlines():
                _append_unique(texts, seen, part)
        elif child.tag.endswith("}tbl"):
            for text in _iter_table_texts(Table(child, doc)):
                _append_unique(texts, seen, text)
    for footer_attr in ("first_page_footer", "footer", "even_page_footer"):
        for section in doc.sections:
            footer = getattr(section, footer_attr)
            for paragraph in footer.paragraphs:
                for part in paragraph.text.splitlines():
                    _append_unique(texts, seen, part)
            for table in footer.tables:
                for text in _iter_table_texts(table):
                    _append_unique(texts, seen, text)
    return texts


def _replace_summary(doc: Document, target: dict[str, Any]) -> None:
    replacement = (
        "Computer Science graduate from Miami University with GPA 3.47 Overall and "
        "3.65 CS Coursework, plus ITSM Data Integration internship experience building "
        "Python/SQL pipelines on 1M+ row datasets, "
        "Tableau dashboards, workflow automations, and full-stack AI/product projects. "
        "Experienced with Python, SQL, TypeScript, React, backend services, data tooling, "
        "machine learning workflows, and human-reviewed AI systems."
    )
    for idx, paragraph in enumerate(doc.paragraphs[:-1]):
        if paragraph.text.strip() == "Summary":
            target_paragraph = doc.paragraphs[idx + 1]
            target_paragraph.clear()
            target_paragraph.add_run(replacement)
            return
    raise ValueError("Could not find Summary section in source resume")


def _replace_gpa_line(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "GPA" in text and ("3.46" in text or "Dept GPA" in text or "Major GPA" in text):
            lines = []
            replaced = False
            for line in paragraph.text.splitlines() or [paragraph.text]:
                if "GPA" in line and ("3.46" in line or "Dept GPA" in line or "Major GPA" in line):
                    lines.append(APPROVED_GPA_LINE)
                    replaced = True
                else:
                    lines.append(line)
            paragraph.clear()
            paragraph.add_run("\n".join(lines) if replaced else APPROVED_GPA_LINE)
            return


SUPPORTED_SUMMARY_TERMS = {
    "python": "Python",
    "sql": "SQL",
    "typescript": "TypeScript",
    "react": "React",
    "backend": "backend services",
    "frontend": "frontend products",
    "data": "data pipelines",
    "ml": "machine learning workflows",
    "ai": "AI workflow systems",
    "apis": "API integrations",
    "platform": "platform tooling",
}


def _source_supported_terms(source: ResumeSource, keywords: list[str]) -> list[str]:
    source_text = source.full_text.lower()
    selected: list[str] = []
    for keyword in keywords:
        normalized = str(keyword).lower().replace("-", " ")
        for needle, label in SUPPORTED_SUMMARY_TERMS.items():
            if needle in normalized and needle in source_text and label not in selected:
                selected.append(label)
    return selected[:6]


def _summary_from_source(source: ResumeSource, target: dict[str, Any]) -> str:
    source_summary = " ".join(source.sections.get("Summary", [])).strip()
    if not source_summary:
        source_summary = (
            "Computer Science graduate from Miami University with experience building "
            "Python and SQL data pipelines, full-stack products, applied ML workflows, "
            "and reliability-focused tooling."
        )
    strengths = _source_supported_terms(source, list(target.get("keywords") or []))
    if strengths:
        return f"{source_summary} Strengths include {', '.join(strengths)}."
    return source_summary


def _build_docx_from_source(source: ResumeSource, target: dict[str, Any]) -> Document:
    doc = Document()
    summary_paragraphs = set(source.sections.get("Summary", []))
    for paragraph in source.paragraphs:
        if paragraph in SECTION_HEADINGS:
            doc.add_paragraph(paragraph)
            if paragraph == "Summary":
                doc.add_paragraph(_summary_from_source(source, target))
            continue
        if summary_paragraphs and paragraph in summary_paragraphs:
            continue
        doc.add_paragraph(paragraph)
    return doc


def _transcript_profile_for_source(source_path: Path) -> Path | None:
    sibling_profile = source_path.parent / "transcript-profile.json"
    if sibling_profile.exists():
        return sibling_profile
    if source_path.resolve() == MAIN_RESUME_DOCX.resolve() and TRANSCRIPT_PROFILE_JSON.exists():
        return TRANSCRIPT_PROFILE_JSON
    return None


def _write_evidence(path: Path, target: dict[str, Any], source_path: Path, evidence: list[dict[str, Any]]) -> None:
    transcript_profile = _transcript_profile_for_source(source_path)
    payload = {
        "target": target,
        "source_resume": str(source_path),
        "transcript_profile": str(transcript_profile) if transcript_profile else "",
        "claims": evidence,
        "review_status": "needs_user_review",
        "rules": [
            "No invented skills, metrics, citizenship, clearance, dates, or outcomes.",
            "Original resume remains unchanged.",
            "Tailored outputs require user review before application use.",
        ],
    }
    path.write_text(json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8")


def _wrap_pdf_text(c: canvas.Canvas, text: str, font: str, size: float, max_width: float) -> list[str]:
    words = text.split()
    if not words:
        return []
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        trial = f"{current} {word}"
        if c.stringWidth(trial, font, size) <= max_width:
            current = trial
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _build_pdf_from_docx(docx_path: Path, pdf_path: Path) -> None:
    doc = Document(docx_path)
    paragraphs = _paragraph_texts(doc)
    width, height = letter

    for body_size in (11.8, 11.6, 11.4, 11.2, 11.0, 10.8, 10.6, 10.4, 10.2, 10.0, 9.8, 9.6, 9.4, 9.2, 9.0, 8.8, 8.4, 8.0, 7.6, 7.2, 6.9):
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        left = 0.5 * inch
        right = 0.5 * inch
        top = height - 0.35 * inch
        y = top
        max_width = width - left - right
        line_height = body_size + 1.25
        overflow = False

        for idx, text in enumerate(paragraphs):
            is_name = idx == 0
            is_contact = idx == 1
            is_heading = text in SECTION_HEADINGS
            if is_name:
                font, size = "Times-Bold", 15
                text_width = c.stringWidth(text, font, size)
                c.setFillColorRGB(0, 0, 0)
                c.setFont(font, size)
                c.drawString((width - text_width) / 2, y, text)
                y -= size + 2
                continue
            if is_contact:
                font, size = "Times-Bold", 8.5
                text_width = c.stringWidth(text, font, size)
                c.setFillColorRGB(0, 0, 0)
                c.setFont(font, size)
                c.drawString(max(left, (width - text_width) / 2), y, text)
                y -= size + 7
                continue
            if is_heading:
                y -= 1
                c.setFillColorRGB(0.12, 0.32, 0.58)
                c.setFont("Times-Bold", body_size + 3)
                c.drawString(left, y, text)
                y -= body_size + 4
                continue

            indent = 12 if text.startswith(("CSE ", "RAG ", "Automated ", "TS ", "On-device", "Vision ", "Wearable", "Desktop", "Tauri", "Future", "Collaborated", "CLI", "Catch2", "Integrated", "Built Python", "Delivered", "Built Workday", "Supported")) else 0
            bullet = indent > 0 or text.startswith(("Agentic", "Visual", "LifeQuest", "Fast", "Dynamic", "ITSM", "Student Worker", "Volunteer", "Finalist", "4th Place"))
            c.setFillColorRGB(0, 0, 0)
            c.setFont("Times-Bold" if bullet and indent == 0 else "Times-Roman", body_size)
            prefix_width = 10 if bullet else 0
            lines = _wrap_pdf_text(c, text, "Times-Bold" if bullet and indent == 0 else "Times-Roman", body_size, max_width - indent - prefix_width)
            if bullet:
                c.setFont("Times-Roman", body_size)
                c.drawString(left + indent, y, "•")
            for line_idx, line in enumerate(lines):
                c.setFont("Times-Bold" if bullet and indent == 0 else "Times-Roman", body_size)
                x = left + indent + (prefix_width if bullet else 0)
                c.drawString(x, y, line)
                y -= line_height
                if y < 0.35 * inch:
                    overflow = True
                    break
            if overflow:
                break

        if not overflow:
            c.save()
            return
        c.save()


def build_pdf_from_docx(docx_path: Path, pdf_path: Path) -> None:
    _build_pdf_from_docx(docx_path, pdf_path)


def _normalized_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("•", "-")).strip()


def _target_forbidden_terms(evidence_path: Path, source_text: str) -> list[str]:
    if not evidence_path.exists():
        return []
    try:
        payload = json.loads(evidence_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []

    target = payload.get("target") if isinstance(payload, dict) else None
    if not isinstance(target, dict):
        return []

    source_lower = source_text.lower()
    terms: list[str] = []
    for key in ("company", "title"):
        value = str(target.get(key) or "").strip()
        if value and value.lower() not in source_lower and value not in terms:
            terms.append(value)
    return terms


def _forbidden_visible_text(pdf_text: str, evidence_path: Path, source_text: str) -> list[str]:
    candidates = [*STATIC_FORBIDDEN_VISIBLE_TEXT, *_target_forbidden_terms(evidence_path, source_text)]
    haystack = pdf_text.lower()
    return [candidate for candidate in candidates if candidate and candidate.lower() in haystack]


def _forbidden_docx_text(texts: list[str], evidence_path: Path, source_text: str) -> list[str]:
    candidates = [*STATIC_FORBIDDEN_VISIBLE_TEXT, *_target_forbidden_terms(evidence_path, source_text)]
    haystack = "\n".join(texts).lower()
    return [candidate for candidate in candidates if candidate and candidate.lower() in haystack]


def _pdf_contains_all_docx_text(texts: list[str], pdf_text: str) -> bool:
    normalized_pdf = _normalized_text(pdf_text)
    for text in texts:
        normalized = _normalized_text(text)
        if normalized and normalized not in normalized_pdf:
            return False
    return True


def create_tailored_resume(
    source_docx: Path,
    output_dir: Path,
    target: dict[str, Any],
    evidence: list[dict[str, Any]],
) -> ResumeArtifacts:
    output_dir.mkdir(parents=True, exist_ok=True)
    company_label = _artifact_company(str(target["company"]))
    docx_path = output_dir / f"Ayush-Yadav-Resume-{{{company_label}}}.docx"
    pdf_path = output_dir / f"Ayush-Yadav-Resume-{{{company_label}}}.pdf"
    evidence_path = output_dir / "evidence-ledger.json"
    validation_path = output_dir / "validation-report.json"

    source = load_resume_source(source_docx)
    if source.source_format == "docx":
        doc = Document(str(source_docx))
        _replace_summary(doc, target)
        _replace_gpa_line(doc)
    else:
        doc = _build_docx_from_source(source, target)
    doc.save(docx_path)
    _write_evidence(evidence_path, target, source_docx, evidence)
    _build_pdf_from_docx(docx_path, pdf_path)

    return ResumeArtifacts(
        docx_path=docx_path,
        pdf_path=pdf_path,
        evidence_path=evidence_path,
        validation_path=validation_path,
    )


def _xml_counts(docx_path: Path) -> dict[str, int]:
    with zipfile.ZipFile(docx_path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", "ignore")
    return {
        "tables": xml.count("<w:tbl"),
        "drawings": xml.count("<w:drawing"),
        "text_boxes": xml.count("<w:txbxContent"),
    }


def _docx_page_size(doc: Document) -> str:
    section = doc.sections[0]
    width = section.page_width.inches
    height = section.page_height.inches
    if abs(width - 8.5) < 0.1 and abs(height - 11.0) < 0.1:
        return "Letter"
    return f"{width:.2f}x{height:.2f}"


def _pdf_page_size(reader: PdfReader) -> str:
    page = reader.pages[0]
    width = float(page.mediabox.width)
    height = float(page.mediabox.height)
    if abs(width - 612) < 2 and abs(height - 792) < 2:
        return "Letter"
    return f"{width:.0f}x{height:.0f}"


def validate_resume_artifacts(artifacts: ResumeArtifacts, source_docx: Path) -> dict[str, Any]:
    doc = Document(artifacts.docx_path)
    source = load_resume_source(source_docx)
    texts = _all_docx_texts(doc)
    full_text = "\n".join(texts)
    source_text = source.full_text
    xml = _xml_counts(artifacts.docx_path)
    reader = PdfReader(str(artifacts.pdf_path))
    page = reader.pages[0]
    cropbox_matches_mediabox = (
        abs(float(page.cropbox.width) - float(page.mediabox.width)) < 1
        and abs(float(page.cropbox.height) - float(page.mediabox.height)) < 1
    )
    pdf_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    ratio = len(full_text) / max(source.text_length, 1)
    pdf_text_ratio_to_docx = len(pdf_text) / max(len(full_text), 1)
    docx_page_size = _docx_page_size(doc)
    page_size = _pdf_page_size(reader)
    pdf_contains_all_docx_text = _pdf_contains_all_docx_text(texts, pdf_text)
    forbidden_visible_text = _forbidden_visible_text(pdf_text, artifacts.evidence_path, source_text)
    forbidden_docx_text = _forbidden_docx_text(texts, artifacts.evidence_path, source_text)
    full_page_estimate = (
        len(reader.pages) == 1
        and ratio >= 0.9
        and len(pdf_text) >= 3000
        and pdf_text_ratio_to_docx >= 0.95
        and pdf_contains_all_docx_text
    )
    has_expected_headings = all(heading in texts for heading in SECTION_HEADINGS)
    is_valid = (
        docx_page_size == "Letter"
        and len(reader.pages) == 1
        and page_size == "Letter"
        and cropbox_matches_mediabox
        and xml["tables"] == 0
        and xml["drawings"] == 0
        and xml["text_boxes"] == 0
        and len(texts) >= 40
        and has_expected_headings
        and ratio >= 0.9
        and pdf_text.startswith("Ayush Yadav\n")
        and full_page_estimate
        and not forbidden_visible_text
        and not forbidden_docx_text
    )
    report = {
        **xml,
        "cropbox_matches_mediabox": cropbox_matches_mediabox,
        "docx_page_size": docx_page_size,
        "page_count": len(reader.pages),
        "page_size": page_size,
        "paragraph_count": len(texts),
        "headings": [text for text in texts if text in SECTION_HEADINGS],
        "text_length": len(full_text),
        "source_text_length": len(source_text),
        "text_length_ratio_to_source": ratio,
        "pdf_text_length_ratio_to_docx": pdf_text_ratio_to_docx,
        "pdf_contains_all_docx_text": pdf_contains_all_docx_text,
        "forbidden_visible_text": forbidden_visible_text,
        "forbidden_docx_text": forbidden_docx_text,
        "pdf_text": pdf_text,
        "full_page_estimate": full_page_estimate,
        "is_valid": is_valid,
    }
    artifacts.validation_path.write_text(json.dumps({k: v for k, v in report.items() if k != "pdf_text"}, indent=2, sort_keys=True), encoding="utf-8")
    return report
