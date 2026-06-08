from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from pypdf.errors import PdfReadError


CANONICAL_HEADINGS = {
    "summary": "Summary",
    "education": "Education",
    "technical skills": "Technical Skills",
    "projects": "Projects",
    "experience": "Experience",
    "leadership & activities": "Leadership & Activities",
    "leadership and activities": "Leadership & Activities",
}

SECTION_ORDER = [
    "Summary",
    "Education",
    "Technical Skills",
    "Projects",
    "Experience",
    "Leadership & Activities",
]


@dataclass(frozen=True)
class ResumeSource:
    path: Path
    source_format: str
    paragraphs: list[str]
    sections: dict[str, list[str]]
    full_text: str
    text_length: int
    page_count: int | None
    page_size: str | None
    transcript_profile: Path | None


def normalize_heading(value: str) -> str:
    stripped = " ".join(value.strip().split())
    return CANONICAL_HEADINGS.get(stripped.lower(), stripped)


def _clean_text(value: str) -> str:
    return " ".join(value.strip().split())


def _page_size(reader: PdfReader) -> str:
    page = reader.pages[0]
    width = float(page.mediabox.width)
    height = float(page.mediabox.height)
    if abs(width - 612) < 2 and abs(height - 792) < 2:
        return "Letter"
    return f"{width:.0f}x{height:.0f}"


def _sections_from_paragraphs(paragraphs: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current: str | None = None
    for paragraph in paragraphs:
        heading = normalize_heading(paragraph)
        if heading in SECTION_ORDER:
            current = heading
            sections.setdefault(current, [])
            continue
        if current:
            sections.setdefault(current, []).append(paragraph)
    return sections


def _profile_for(path: Path) -> Path | None:
    profile = path.parent / "transcript-profile.json"
    return profile if profile.exists() else None


def _append_unique(texts: list[str], seen: set[str], text: str) -> None:
    if text and text not in seen:
        texts.append(text)
        seen.add(text)


def _iter_table_paragraphs(table: Table) -> list[str]:
    texts: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = _clean_text(paragraph.text)
                if text:
                    texts.append(text)
            for nested_table in cell.tables:
                texts.extend(_iter_table_paragraphs(nested_table))
    return texts


def _docx_body_paragraphs(doc: Document) -> list[str]:
    texts: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            text = _clean_text(Paragraph(child, doc).text)
            if text:
                texts.append(text)
        elif child.tag.endswith("}tbl"):
            texts.extend(_iter_table_paragraphs(Table(child, doc)))
    return texts


def _docx_header_footer_paragraphs(doc: Document) -> tuple[list[str], list[str]]:
    header_texts: list[str] = []
    footer_texts: list[str] = []
    header_seen: set[str] = set()
    footer_seen: set[str] = set()
    for section in doc.sections:
        for header in (section.first_page_header, section.header, section.even_page_header):
            for paragraph in header.paragraphs:
                _append_unique(header_texts, header_seen, _clean_text(paragraph.text))
            for table in header.tables:
                for text in _iter_table_paragraphs(table):
                    _append_unique(header_texts, header_seen, text)
        for footer in (section.first_page_footer, section.footer, section.even_page_footer):
            for paragraph in footer.paragraphs:
                _append_unique(footer_texts, footer_seen, _clean_text(paragraph.text))
            for table in footer.tables:
                for text in _iter_table_paragraphs(table):
                    _append_unique(footer_texts, footer_seen, text)
    return header_texts, footer_texts


def _docx_paragraphs(doc: Document) -> tuple[list[str], list[str]]:
    body_texts = _docx_body_paragraphs(doc)
    header_texts, footer_texts = _docx_header_footer_paragraphs(doc)
    return [*header_texts, *body_texts, *footer_texts], body_texts


def _load_docx(path: Path) -> ResumeSource:
    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"Could not read DOCX resume source {path}: {exc}") from exc

    paragraphs, section_paragraphs = _docx_paragraphs(doc)
    if not paragraphs:
        raise ValueError(f"DOCX resume source has no extractable text: {path}")

    full_text = "\n".join(paragraphs)
    return ResumeSource(
        path=path,
        source_format="docx",
        paragraphs=paragraphs,
        sections=_sections_from_paragraphs(section_paragraphs),
        full_text=full_text,
        text_length=len(full_text),
        page_count=None,
        page_size=None,
        transcript_profile=_profile_for(path),
    )


def _pdf_reader(path: Path) -> PdfReader:
    try:
        reader = PdfReader(str(path))
    except (PdfReadError, OSError, ValueError) as exc:
        raise ValueError(f"Could not read PDF resume source {path}: {exc}") from exc

    if reader.is_encrypted:
        raise ValueError(f"PDF resume source is encrypted: {path}")
    if len(reader.pages) == 0:
        raise ValueError(f"PDF resume source has no pages: {path}")
    return reader


def _ends_like_wrapped_line(value: str) -> bool:
    stripped = value.rstrip()
    lowered = stripped.lower()
    return stripped.endswith((",", "-", "/", "&")) or lowered.endswith(
        (" and", " or", " with", " plus", " including", " using")
    )


def _starts_like_continuation(value: str) -> bool:
    stripped = value.lstrip()
    return bool(stripped) and (stripped[0].islower() or stripped[0] in ",;:)]")


def _ends_terminally(value: str) -> bool:
    return value.rstrip().endswith((".", ";", ":"))


def _coalesce_pdf_lines(lines: list[str]) -> list[str]:
    paragraphs: list[str] = []
    current_section: str | None = None

    for line in lines:
        heading = normalize_heading(line)
        if heading in SECTION_ORDER:
            paragraphs.append(heading)
            current_section = heading
            continue

        if not paragraphs:
            paragraphs.append(line)
            continue

        previous = paragraphs[-1]
        previous_is_heading = previous in SECTION_ORDER
        starts_bullet = line.startswith(("- ", "• "))
        previous_starts_bullet = previous.startswith(("- ", "• "))
        should_join_wrapped_line = (
            previous_starts_bullet
            and not starts_bullet
            and not _ends_terminally(previous)
        ) or (
            current_section == "Summary"
            and not previous_is_heading
            and not starts_bullet
            and not previous_starts_bullet
            and not _ends_terminally(previous)
        ) or (
            current_section in {"Education", "Technical Skills"}
            and not previous_is_heading
            and not starts_bullet
            and not previous_starts_bullet
            and (_ends_like_wrapped_line(previous) or _starts_like_continuation(line))
        )

        if should_join_wrapped_line:
            paragraphs[-1] = f"{previous} {line}"
        else:
            paragraphs.append(line)

    return paragraphs


def _pdf_lines(reader: PdfReader, path: Path) -> list[str]:
    lines: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            raise ValueError(f"Could not extract text from PDF resume source {path} page {page_index}: {exc}") from exc
        for line in text.splitlines():
            stripped = _clean_text(line)
            if stripped:
                lines.append(stripped)
    return lines


def _load_pdf(path: Path) -> ResumeSource:
    reader = _pdf_reader(path)
    paragraphs = _coalesce_pdf_lines(_pdf_lines(reader, path))
    if not paragraphs:
        raise ValueError(f"PDF resume source has no extractable text: {path}")

    full_text = "\n".join(paragraphs)
    return ResumeSource(
        path=path,
        source_format="pdf",
        paragraphs=paragraphs,
        sections=_sections_from_paragraphs(paragraphs),
        full_text=full_text,
        text_length=len(full_text),
        page_count=len(reader.pages),
        page_size=_page_size(reader),
        transcript_profile=_profile_for(path),
    )


def load_resume_source(path: Path) -> ResumeSource:
    if not path.exists():
        raise FileNotFoundError(path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _load_docx(path)
    if suffix == ".pdf":
        return _load_pdf(path)
    raise ValueError(f"Unsupported resume source format: {path}")
