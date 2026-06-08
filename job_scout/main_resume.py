from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from job_scout.resume_pipeline import (
    ResumeArtifacts,
    build_pdf_from_docx,
    validate_resume_artifacts,
)
from job_scout.transcript_profile import write_transcript_profile


@dataclass(frozen=True)
class MainResumeArtifacts:
    docx_path: Path
    pdf_path: Path
    transcript_profile_path: Path
    validation_path: Path


def _set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def _replace_after_heading(doc: Document, heading: str, replacement: str) -> None:
    for idx, paragraph in enumerate(doc.paragraphs[:-1]):
        if paragraph.text.strip() == heading:
            _set_paragraph_text(doc.paragraphs[idx + 1], replacement)
            return
    raise ValueError(f"Could not find {heading} section in source resume")


def _replace_matching_paragraph(doc: Document, startswith: str, replacement: str) -> None:
    for paragraph in doc.paragraphs:
        lines = paragraph.text.splitlines()
        for line_idx, line in enumerate(lines):
            if line.strip().startswith(startswith):
                lines[line_idx] = replacement
                _set_paragraph_text(paragraph, "\n".join(lines))
                return
    raise ValueError(f"Could not find paragraph starting with {startswith!r}")


def create_main_resume(source_docx: Path, output_dir: Path) -> MainResumeArtifacts:
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = output_dir / "Ayush-Yadav-Main-Resume-2026.docx"
    pdf_path = output_dir / "Ayush-Yadav-Main-Resume-2026.pdf"
    transcript_profile_path = output_dir / "transcript-profile.json"
    validation_path = output_dir / "resume-validation-report.json"

    profile = write_transcript_profile(transcript_profile_path)
    doc = Document(str(source_docx))
    summary = (
        "Computer Science graduate from Miami University with strengths in software engineering, "
        "data systems, machine learning, algorithms, and high-performance computing. Built "
        "Python/SQL pipelines on 1M+ row datasets, Tableau dashboards, workflow automations, "
        "and full-stack AI/product projects using Python, Java, C++, TypeScript, React, SQL, "
        "PyTorch, and OpenMP."
    )
    _replace_after_heading(doc, "Summary", summary)
    _replace_matching_paragraph(
        doc,
        "Miami University, Oxford, OH",
        "Miami University, Oxford, OH | Bachelor of Science in Computer Science, May 2026",
    )
    _replace_matching_paragraph(doc, "GPA:", profile["gpa_labels"]["resume_default"])
    _replace_matching_paragraph(
        doc,
        "CSE 385",
        "CSE 385 - Database Systems; CSE 432 - Machine Learning; "
        "CSE 484 - Algorithms II; CSE 443 - High Performance Computing",
    )
    _replace_matching_paragraph(
        doc,
        "CSE 432",
        "CSE 212 - Software Engineering for UI/UX; "
        "CSE 448/449 - Senior Design Project I/II",
    )
    for marker in ("CSE 443", "CSE 484"):
        for paragraph in doc.paragraphs:
            if paragraph.text.strip().startswith(marker):
                _delete_paragraph(paragraph)

    doc.save(docx_path)
    build_pdf_from_docx(docx_path, pdf_path)
    resume_artifacts = ResumeArtifacts(
        docx_path=docx_path,
        pdf_path=pdf_path,
        evidence_path=transcript_profile_path,
        validation_path=validation_path,
    )
    report = validate_resume_artifacts(resume_artifacts, source_docx)
    validation_path.write_text(
        json.dumps({k: v for k, v in report.items() if k != "pdf_text"}, indent=2, sort_keys=True),
        encoding="utf-8",
    )
    return MainResumeArtifacts(
        docx_path=docx_path,
        pdf_path=pdf_path,
        transcript_profile_path=transcript_profile_path,
        validation_path=validation_path,
    )
