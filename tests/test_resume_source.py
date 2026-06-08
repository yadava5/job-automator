import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from job_scout.config import NEW_RESUME_PDF
from job_scout.resume_source import load_resume_source, normalize_heading
from pypdf import PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def _write_pdf(path: Path, lines: list[str]) -> None:
    pdf = canvas.Canvas(str(path), pagesize=letter)
    y = 760
    for line in lines:
        pdf.drawString(36, y, line)
        y -= 14
    pdf.save()


class ResumeSourceTests(unittest.TestCase):
    def test_normalize_heading_accepts_resume_variants(self):
        self.assertEqual(normalize_heading("Leadership and Activities"), "Leadership & Activities")
        self.assertEqual(normalize_heading("Leadership & Activities"), "Leadership & Activities")
        self.assertEqual(normalize_heading("Technical Skills"), "Technical Skills")

    def test_load_generated_pdf_source_extracts_sections_and_wrapped_lines(self):
        with TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / "fixture.pdf"
            _write_pdf(
                pdf_path,
                [
                    "Ayush Yadav",
                    "ayush@example.com",
                    "Summary",
                    "Computer Science graduate with Python, SQL,",
                    "React, TypeScript, and ML project experience.",
                    "Education",
                    "Miami University, B.S. Computer Science",
                    "GPA: 3.47 Overall",
                    "Technical Skills",
                    "Languages: Python, SQL,",
                    "TypeScript, JavaScript",
                    "Projects",
                    "- Built data pipeline with",
                    "1M+ row validations.",
                    "Experience",
                    "- Delivered dashboards with",
                    "Tableau and SQL.",
                    "Leadership and Activities",
                    "Dean's List",
                ],
            )

            source = load_resume_source(pdf_path)

        self.assertEqual(source.path, pdf_path)
        self.assertEqual(source.source_format, "pdf")
        self.assertEqual(source.page_count, 1)
        self.assertEqual(source.page_size, "Letter")
        self.assertIn("Summary", source.sections)
        self.assertIn("Education", source.sections)
        self.assertIn("Technical Skills", source.sections)
        self.assertIn("Projects", source.sections)
        self.assertIn("Experience", source.sections)
        self.assertIn("Leadership & Activities", source.sections)
        self.assertIn(
            "Computer Science graduate with Python, SQL, React, TypeScript, and ML project experience.",
            source.sections["Summary"],
        )
        self.assertIn("Languages: Python, SQL, TypeScript, JavaScript", source.sections["Technical Skills"])
        self.assertIn("- Built data pipeline with 1M+ row validations.", source.sections["Projects"])
        self.assertIn("- Delivered dashboards with Tableau and SQL.", source.sections["Experience"])

    def test_load_pdf_keeps_distinct_education_and_skill_rows(self):
        with TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / "fixture.pdf"
            _write_pdf(
                pdf_path,
                [
                    "Ayush Yadav",
                    "Summary",
                    "Computer Science graduate.",
                    "Education",
                    "Miami University, B.S. Computer Science",
                    "GPA: 3.47 Overall",
                    "Technical Skills",
                    "Languages: Python, SQL,",
                    "TypeScript, JavaScript",
                    "Frameworks: React, Node.js",
                    "Projects",
                    "- Built data pipeline with",
                    "1M+ row validations.",
                ],
            )

            source = load_resume_source(pdf_path)

        self.assertEqual(
            source.sections["Education"],
            ["Miami University, B.S. Computer Science", "GPA: 3.47 Overall"],
        )
        self.assertEqual(
            source.sections["Technical Skills"],
            ["Languages: Python, SQL, TypeScript, JavaScript", "Frameworks: React, Node.js"],
        )

    @unittest.skipUnless(NEW_RESUME_PDF.exists(), "requires local synced resume PDF")
    def test_load_local_new_resume_pdf_source_contract(self):
        source = load_resume_source(NEW_RESUME_PDF)

        self.assertEqual(source.path, NEW_RESUME_PDF)
        self.assertEqual(source.source_format, "pdf")
        self.assertEqual(source.page_count, 1)
        self.assertEqual(source.page_size, "Letter")
        self.assertGreaterEqual(source.text_length, 3800)
        self.assertTrue(source.full_text.startswith("Ayush Yadav"))
        self.assertIn("Summary", source.sections)
        self.assertIn("Education", source.sections)
        self.assertIn("Technical Skills", source.sections)
        self.assertIn("Projects", source.sections)
        self.assertIn("Experience", source.sections)
        self.assertIn("Leadership & Activities", source.sections)
        self.assertIn("3.47 Overall", source.full_text)
        self.assertNotIn("Computer Science undergraduate", source.full_text)
        self.assertNotIn("transcript-backed strengths", source.full_text)

    def test_load_missing_source_fails_clearly(self):
        missing = Path("/tmp/ayush-missing-resume.pdf")

        with self.assertRaises(FileNotFoundError):
            load_resume_source(missing)

    def test_load_empty_pdf_fails_clearly(self):
        with TemporaryDirectory() as tmp:
            empty_pdf = Path(tmp) / "empty.pdf"
            writer = PdfWriter()
            with empty_pdf.open("wb") as handle:
                writer.write(handle)

            with self.assertRaisesRegex(ValueError, "no pages"):
                load_resume_source(empty_pdf)

    def test_load_encrypted_pdf_fails_clearly(self):
        with TemporaryDirectory() as tmp:
            encrypted_pdf = Path(tmp) / "encrypted.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=612, height=792)
            writer.encrypt("secret")
            with encrypted_pdf.open("wb") as handle:
                writer.write(handle)

            with self.assertRaisesRegex(ValueError, "encrypted"):
                load_resume_source(encrypted_pdf)

    def test_load_blank_pdf_fails_clearly(self):
        with TemporaryDirectory() as tmp:
            blank_pdf = Path(tmp) / "blank.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=612, height=792)
            with blank_pdf.open("wb") as handle:
                writer.write(handle)

            with self.assertRaisesRegex(ValueError, "no extractable text"):
                load_resume_source(blank_pdf)

    def test_load_malformed_pdf_fails_clearly(self):
        with TemporaryDirectory() as tmp:
            malformed_pdf = Path(tmp) / "malformed.pdf"
            malformed_pdf.write_text("not a pdf", encoding="utf-8")

            with self.assertRaisesRegex(ValueError, "Could not read PDF resume source"):
                load_resume_source(malformed_pdf)

    def test_load_docx_includes_table_header_and_footer_text(self):
        with TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "fixture.docx"
            doc = Document()
            doc.sections[0].header.paragraphs[0].text = "Header Contact"
            doc.sections[0].footer.paragraphs[0].text = "Footer Portfolio"
            doc.add_paragraph("Ayush Yadav")
            doc.add_paragraph("Summary")
            doc.add_paragraph("Computer Science graduate.")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Technical Skills"
            table.cell(0, 1).text = "Python, SQL, TypeScript"
            doc.save(docx_path)

            source = load_resume_source(docx_path)

        self.assertIn("Header Contact", source.paragraphs)
        self.assertIn("Footer Portfolio", source.paragraphs)
        self.assertIn("Technical Skills", source.sections)
        self.assertIn("Python, SQL, TypeScript", source.sections["Technical Skills"])
        self.assertNotIn("Header Contact", source.sections["Technical Skills"])
        self.assertNotIn("Footer Portfolio", source.sections["Technical Skills"])

    def test_load_docx_includes_first_page_and_even_page_header_footer_text(self):
        with TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "fixture.docx"
            doc = Document()
            section = doc.sections[0]
            section.different_first_page_header_footer = True
            section.odd_and_even_pages_header_footer = True
            section.first_page_header.paragraphs[0].text = "First Page Contact"
            section.first_page_footer.paragraphs[0].text = "First Page Portfolio"
            section.even_page_header.paragraphs[0].text = "Even Page Contact"
            section.even_page_footer.paragraphs[0].text = "Even Page Portfolio"
            doc.add_paragraph("Ayush Yadav")
            doc.add_paragraph("Summary")
            doc.add_paragraph("Computer Science graduate.")
            doc.save(docx_path)

            source = load_resume_source(docx_path)

        self.assertIn("First Page Contact", source.paragraphs)
        self.assertIn("First Page Portfolio", source.paragraphs)
        self.assertIn("Even Page Contact", source.paragraphs)
        self.assertIn("Even Page Portfolio", source.paragraphs)
        self.assertNotIn("First Page Contact", source.sections["Summary"])
        self.assertNotIn("Even Page Portfolio", source.sections["Summary"])


if __name__ == "__main__":
    unittest.main()
