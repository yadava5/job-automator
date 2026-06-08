import json
import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document

from job_scout.config import MAIN_RESUME_DOCX, NEW_RESUME_PDF, ORIGINAL_RESUME_DOCX
from job_scout.resume_pipeline import build_pdf_from_docx, create_tailored_resume, validate_resume_artifacts


SOURCE_RESUME = ORIGINAL_RESUME_DOCX


class ResumePipelineTests(unittest.TestCase):
    def test_tailored_resume_preserves_original_structure_and_validates_pdf(self):
        target = {
            "company": "Instabase",
            "title": "Full-stack Software Engineer (New Grad)",
            "keywords": ["full-stack", "REST APIs", "AI innovation", "backend", "frontend"],
        }
        evidence = [
            {
                "claim": "Full-stack AI workflow platform",
                "evidence_source": "original_resume",
                "evidence_path": str(SOURCE_RESUME),
                "confidence": "verified",
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            artifacts = create_tailored_resume(SOURCE_RESUME, Path(tmp), target, evidence)
            report = validate_resume_artifacts(artifacts, SOURCE_RESUME)
            self.assertTrue(artifacts.docx_path.exists())
            self.assertTrue(artifacts.pdf_path.exists())
            self.assertTrue(artifacts.evidence_path.exists())
            self.assertEqual(report["page_count"], 1)
            self.assertEqual(report["page_size"], "Letter")
            self.assertEqual(report["tables"], 0)
            self.assertEqual(report["drawings"], 0)
            self.assertEqual(report["text_boxes"], 0)
            self.assertGreaterEqual(report["paragraph_count"], 40)
            self.assertGreaterEqual(report["text_length_ratio_to_source"], 0.9)
            self.assertIn("Summary", report["headings"])
            self.assertIn("Technical Skills", report["headings"])
            self.assertTrue(report["pdf_text"].startswith("Ayush Yadav\n"))
            self.assertNotIn("Tailored for", report["pdf_text"])
            self.assertNotIn("Instabase", report["pdf_text"])
            self.assertNotIn("GPA 3.46", report["pdf_text"])
            self.assertNotIn("Dept GPA", report["pdf_text"])
            self.assertNotIn("Major GPA", report["pdf_text"])
            self.assertNotIn("Computer Science undergraduate", report["pdf_text"])
            self.assertIn("GPA 3.47 Overall", report["pdf_text"])
            self.assertTrue(report["full_page_estimate"])
            self.assertTrue(report["is_valid"])


class ResumePipelineGraduateSourceTests(unittest.TestCase):
    def test_tailored_resume_uses_graduate_wording_from_main_resume(self):
        target = {
            "company": "Instabase",
            "title": "Full-stack Software Engineer (New Grad)",
            "keywords": ["full-stack", "REST APIs", "AI innovation", "backend", "frontend"],
        }
        evidence = [
            {
                "claim": "Verified main resume",
                "evidence_source": "main_resume",
                "evidence_path": str(MAIN_RESUME_DOCX),
                "confidence": "verified",
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            artifacts = create_tailored_resume(MAIN_RESUME_DOCX, Path(tmp), target, evidence)
            report = validate_resume_artifacts(artifacts, MAIN_RESUME_DOCX)
            ledger = json.loads(artifacts.evidence_path.read_text(encoding="utf-8"))

        self.assertIn("Computer Science graduate", report["pdf_text"])
        self.assertNotIn("Computer Science undergraduate", report["pdf_text"])
        self.assertNotIn("Tailored for", report["pdf_text"])
        self.assertNotIn("Instabase", report["pdf_text"])
        self.assertNotIn("Full-stack Software Engineer", report["pdf_text"])
        expected_profile = MAIN_RESUME_DOCX.parent / "transcript-profile.json"
        self.assertEqual(ledger["transcript_profile"], str(expected_profile))
        self.assertTrue(Path(ledger["transcript_profile"]).exists())

    def test_evidence_ledger_uses_source_sibling_transcript_profile(self):
        target = {
            "company": "Instabase",
            "title": "Full-stack Software Engineer (New Grad)",
            "keywords": ["full-stack", "backend"],
        }
        evidence = [
            {
                "claim": "Alternate verified main resume",
                "evidence_source": "alternate_main_resume",
                "evidence_path": str(MAIN_RESUME_DOCX),
                "confidence": "verified",
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_dir = tmp_path / "alternate-main-resume"
            source_dir.mkdir()
            source_docx = source_dir / "Alternate-Main-Resume.docx"
            shutil.copy2(MAIN_RESUME_DOCX, source_docx)
            expected_profile = source_dir / "transcript-profile.json"
            expected_profile.write_text('{"source": "alternate"}\n', encoding="utf-8")

            artifacts = create_tailored_resume(source_docx, tmp_path / "pack", target, evidence)
            ledger = json.loads(artifacts.evidence_path.read_text(encoding="utf-8"))

            self.assertEqual(ledger["source_resume"], str(source_docx))
            self.assertEqual(ledger["transcript_profile"], str(expected_profile))
            self.assertTrue(Path(ledger["transcript_profile"]).exists())

    def test_evidence_ledger_omits_profile_when_source_has_no_sibling_profile(self):
        target = {
            "company": "Instabase",
            "title": "Full-stack Software Engineer (New Grad)",
            "keywords": ["full-stack", "backend"],
        }
        evidence = [
            {
                "claim": "Alternate main resume without transcript profile",
                "evidence_source": "alternate_main_resume",
                "evidence_path": str(MAIN_RESUME_DOCX),
                "confidence": "verified",
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_dir = tmp_path / "alternate-main-resume"
            source_dir.mkdir()
            source_docx = source_dir / "Alternate-Main-Resume.docx"
            shutil.copy2(MAIN_RESUME_DOCX, source_docx)

            artifacts = create_tailored_resume(source_docx, tmp_path / "pack", target, evidence)
            ledger = json.loads(artifacts.evidence_path.read_text(encoding="utf-8"))

        self.assertEqual(ledger["source_resume"], str(source_docx))
        self.assertEqual(ledger["transcript_profile"], "")


@unittest.skipUnless(NEW_RESUME_PDF.exists(), "requires local synced resume PDF")
class ResumePipelineNewPdfSourceTests(unittest.TestCase):
    def test_application_resume_uses_new_pdf_source_without_old_resume_leakage(self):
        target = {
            "company": "Giga",
            "title": "Software Engineer (New Grads)",
            "keywords": ["backend", "AI", "Python", "React"],
        }
        evidence = [
            {
                "claim": "New iCloud resume source",
                "evidence_source": "new_resume_pdf",
                "evidence_path": str(NEW_RESUME_PDF),
                "confidence": "verified",
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            artifacts = create_tailored_resume(NEW_RESUME_PDF, Path(tmp), target, evidence)
            report = validate_resume_artifacts(artifacts, NEW_RESUME_PDF)
            ledger = json.loads(artifacts.evidence_path.read_text(encoding="utf-8"))
            self.assertTrue(artifacts.docx_path.exists())
            self.assertTrue(artifacts.pdf_path.exists())

        self.assertEqual(artifacts.docx_path.name, "Ayush-Yadav-Resume-{Giga}.docx")
        self.assertEqual(artifacts.pdf_path.name, "Ayush-Yadav-Resume-{Giga}.pdf")
        self.assertEqual(ledger["source_resume"], str(NEW_RESUME_PDF))
        self.assertEqual(report["page_count"], 1)
        self.assertEqual(report["page_size"], "Letter")
        self.assertTrue(report["full_page_estimate"])
        self.assertTrue(report["is_valid"])
        self.assertGreaterEqual(report["text_length"], 3400)
        self.assertGreaterEqual(report["text_length_ratio_to_source"], 0.86)
        for heading in ("Summary", "Education", "Technical Skills", "Projects", "Experience", "Leadership & Activities"):
            self.assertIn(heading, report["headings"])
        for forbidden in (
            "Tailored for",
            "Giga",
            "Software Engineer (New Grads)",
            "transcript-backed strengths",
            "Computer Science undergraduate",
            "GPA 3.46",
            "Dept GPA",
            "Major GPA",
            "Ayush Yadav Resume (Dec 2025)",
        ):
            self.assertNotIn(forbidden, report["pdf_text"])
        self.assertIn("Computer Science graduate", report["pdf_text"])
        self.assertIn("GPA: 3.47 Overall", report["pdf_text"])

    def test_validation_fails_when_generated_heading_is_missing(self):
        target = {
            "company": "Giga",
            "title": "Software Engineer (New Grads)",
            "keywords": ["backend"],
        }
        evidence = [
            {
                "claim": "New iCloud resume source",
                "evidence_source": "new_resume_pdf",
                "evidence_path": str(NEW_RESUME_PDF),
                "confidence": "verified",
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            artifacts = create_tailored_resume(NEW_RESUME_PDF, Path(tmp), target, evidence)
            doc = Document(str(artifacts.docx_path))
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() == "Projects":
                    paragraph.clear()
                    paragraph.add_run("Selected Work")
                    break
            doc.save(artifacts.docx_path)
            build_pdf_from_docx(artifacts.docx_path, artifacts.pdf_path)
            report = validate_resume_artifacts(artifacts, NEW_RESUME_PDF)

        self.assertFalse(report["is_valid"])
        self.assertNotIn("Projects", report["headings"])

    def test_validation_fails_when_forbidden_visible_label_leaks(self):
        target = {
            "company": "Giga",
            "title": "Software Engineer (New Grads)",
            "keywords": ["backend"],
        }
        evidence = [
            {
                "claim": "New iCloud resume source",
                "evidence_source": "new_resume_pdf",
                "evidence_path": str(NEW_RESUME_PDF),
                "confidence": "verified",
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            artifacts = create_tailored_resume(NEW_RESUME_PDF, Path(tmp), target, evidence)
            doc = Document(str(artifacts.docx_path))
            doc.paragraphs[2].insert_paragraph_before("Tailored for Giga Software Engineer (New Grads)")
            doc.save(artifacts.docx_path)
            build_pdf_from_docx(artifacts.docx_path, artifacts.pdf_path)
            report = validate_resume_artifacts(artifacts, NEW_RESUME_PDF)

        self.assertFalse(report["is_valid"])
        self.assertIn("Tailored for", report["forbidden_visible_text"])
        self.assertIn("Giga", report["forbidden_visible_text"])
        self.assertIn("Software Engineer (New Grads)", report["forbidden_visible_text"])

    def test_validation_fails_when_forbidden_visible_label_leaks_with_case_change(self):
        target = {
            "company": "Giga",
            "title": "Software Engineer (New Grads)",
            "keywords": ["backend"],
        }
        evidence = [
            {
                "claim": "New iCloud resume source",
                "evidence_source": "new_resume_pdf",
                "evidence_path": str(NEW_RESUME_PDF),
                "confidence": "verified",
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            artifacts = create_tailored_resume(NEW_RESUME_PDF, Path(tmp), target, evidence)
            doc = Document(str(artifacts.docx_path))
            doc.paragraphs[2].insert_paragraph_before("tailored for GIGA software engineer (new grads)")
            doc.save(artifacts.docx_path)
            build_pdf_from_docx(artifacts.docx_path, artifacts.pdf_path)
            report = validate_resume_artifacts(artifacts, NEW_RESUME_PDF)

        self.assertFalse(report["is_valid"])
        self.assertIn("Tailored for", report["forbidden_visible_text"])
        self.assertIn("Giga", report["forbidden_visible_text"])
        self.assertIn("Software Engineer (New Grads)", report["forbidden_visible_text"])

    def test_validation_fails_when_pdf_omits_generated_docx_text(self):
        target = {
            "company": "Giga",
            "title": "Software Engineer (New Grads)",
            "keywords": ["backend"],
        }
        evidence = [
            {
                "claim": "New iCloud resume source",
                "evidence_source": "new_resume_pdf",
                "evidence_path": str(NEW_RESUME_PDF),
                "confidence": "verified",
            }
        ]
        sentinel = "SENTINEL_VISIBLE_TRAILING_TEXT_REQUIRED"

        with tempfile.TemporaryDirectory() as tmp:
            artifacts = create_tailored_resume(NEW_RESUME_PDF, Path(tmp), target, evidence)
            doc = Document(str(artifacts.docx_path))
            for index in range(80):
                doc.add_paragraph(f"Overflow validation paragraph {index}")
            doc.add_paragraph(sentinel)
            doc.save(artifacts.docx_path)
            build_pdf_from_docx(artifacts.docx_path, artifacts.pdf_path)
            report = validate_resume_artifacts(artifacts, NEW_RESUME_PDF)

        self.assertFalse(report["is_valid"])
        self.assertFalse(report["pdf_contains_all_docx_text"])
        self.assertNotIn(sentinel, report["pdf_text"])

    def test_validation_fails_when_pdf_omits_generated_header_footer_text(self):
        target = {
            "company": "Giga",
            "title": "Software Engineer (New Grads)",
            "keywords": ["backend"],
        }
        evidence = [
            {
                "claim": "New iCloud resume source",
                "evidence_source": "new_resume_pdf",
                "evidence_path": str(NEW_RESUME_PDF),
                "confidence": "verified",
            }
        ]
        footer_sentinel = "SENTINEL_FOOTER_TEXT_REQUIRED"

        with tempfile.TemporaryDirectory() as tmp:
            artifacts = create_tailored_resume(NEW_RESUME_PDF, Path(tmp), target, evidence)
            doc = Document(str(artifacts.docx_path))
            doc.sections[0].header.paragraphs[0].text = "Tailored for Giga Software Engineer (New Grads)"
            doc.sections[0].footer.paragraphs[0].text = footer_sentinel
            doc.save(artifacts.docx_path)
            build_pdf_from_docx(artifacts.docx_path, artifacts.pdf_path)
            report = validate_resume_artifacts(artifacts, NEW_RESUME_PDF)

        self.assertFalse(report["is_valid"])
        self.assertFalse(report["pdf_contains_all_docx_text"])
        self.assertIn("Tailored for", report["forbidden_docx_text"])
        self.assertIn("Giga", report["forbidden_docx_text"])
        self.assertNotIn(footer_sentinel, report["pdf_text"])


if __name__ == "__main__":
    unittest.main()
