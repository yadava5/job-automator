import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from job_scout.config import ORIGINAL_RESUME_DOCX
from job_scout.main_resume import create_main_resume


class MainResumeTests(unittest.TestCase):
    def test_create_main_resume_updates_education_without_overwriting_original(self):
        original_before = ORIGINAL_RESUME_DOCX.read_bytes()

        with tempfile.TemporaryDirectory() as tmp:
            artifacts = create_main_resume(ORIGINAL_RESUME_DOCX, Path(tmp))
            doc = Document(str(artifacts.docx_path))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            report = json.loads(artifacts.validation_path.read_text(encoding="utf-8"))

        self.assertEqual(ORIGINAL_RESUME_DOCX.read_bytes(), original_before)
        self.assertIn("Bachelor of Science in Computer Science, May 2026", text)
        self.assertIn("GPA: 3.47 Overall | 3.65 CS Coursework", text)
        self.assertIn("with strengths in software engineering, data systems", text)
        self.assertNotIn("transcript-backed", text)
        self.assertIn("Database Systems", text)
        self.assertIn("Algorithms II", text)
        self.assertNotIn("Expected May 2026", text)
        self.assertFalse(
            [
                paragraph.style.name
                for paragraph in doc.paragraphs
                if paragraph.style.name.startswith("List") and not paragraph.text.strip()
            ],
            "main resume should not contain empty bullet paragraphs",
        )
        self.assertTrue(report["is_valid"])

    def test_create_main_resume_outputs_one_letter_pdf(self):
        with tempfile.TemporaryDirectory() as tmp:
            artifacts = create_main_resume(ORIGINAL_RESUME_DOCX, Path(tmp))
            reader = PdfReader(str(artifacts.pdf_path))
            page = reader.pages[0]

        self.assertEqual(len(reader.pages), 1)
        self.assertAlmostEqual(float(page.mediabox.width), 612.0, delta=2)
        self.assertAlmostEqual(float(page.mediabox.height), 792.0, delta=2)
        self.assertAlmostEqual(float(page.cropbox.width), 612.0, delta=2)
        self.assertAlmostEqual(float(page.cropbox.height), 792.0, delta=2)


if __name__ == "__main__":
    unittest.main()
